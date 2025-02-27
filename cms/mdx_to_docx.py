import os
import re
from typing import Dict, Callable

## https://python-docx.readthedocs.io/en/latest/user/install.html
## pip install python-docx
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.shared import qn, OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE

def mdx_to_docx(mdx_directory, output_file):
    doc = Document()
    
    for filename in os.listdir(mdx_directory):
        if filename.endswith('.mdx'):
            with open(os.path.join(mdx_directory, filename), 'r') as file:
                content = file.read()
                
                # Add a section marker
                doc.add_heading(f"FILE: {filename}", level=1)
                
                # Extract front matter
                front_matter = re.search(r'---\n(.*?)\n---', content, re.DOTALL)
                # Extract front matter
                front_matter = re.search(r'---\n(.*?)\n---', content, re.DOTALL)
                if front_matter:
                    front_matter_content = front_matter.group(1)
                    # List of keys to include
                    keys_to_include = ['title', 'description', 'id']
                    
                    # Parse front matter content
                    lines = front_matter_content.split('\n')
                    id_value = None
                    i = 0
                    while i < len(lines):
                        line = lines[i].strip()
                        if ':' in line:
                            key, value = line.split(':', 1)
                            key = key.strip()
                            value = value.strip()
                            
                            # Check if the key is in the list of keys to include
                            if key in keys_to_include:
                                # Handle multi-line values
                                if value == '"':
                                    i += 1
                                    while i < len(lines) and not lines[i].strip().endswith('"'):
                                        value += ' ' + lines[i].strip()
                                        i += 1
                                    if i < len(lines):
                                        value += ' ' + lines[i].strip()
                                
                                # Remove surrounding quotes if present
                                value = value.strip('"')
                                
                                if key == 'id':
                                    id_value = value
                                else:
                                    doc.add_paragraph(f"{key.capitalize()}: {value}")
                        i += 1
                    
                    # Add the link if id was found
                    if id_value is not None:
                        p = doc.add_paragraph()
                        add_hyperlink(p, f"https://earth.gov/stories/{id_value}", f"https://earth.gov/stories/{id_value}")
                
                
                # Remove front matter from the content
                if front_matter:
                    content = content.replace(front_matter.group(0), '', 1).strip()
                
                # Remove lines that start with 'import'
                content = '\n'.join([line for line in content.split('\n') if not line.strip().startswith('import')])
                content = content.strip()

                # Add the rest of the content
                render_content_to_doc(doc, content)
                
                # Add a page break between files
                doc.add_page_break()
    
    doc.save(output_file)

def render_content_to_doc(doc: DocumentType, content: str):
    tag_handlers: Dict[str, Callable] = {
        'Carousel': handle_carousel,
        # Add more tag handlers here as needed
    }

    tags = split_content_into_tags(content)
    for i in range(len(tags)):
        tags[i] = split_tag_into_subtags(tags[i])
    
    doc.add_paragraph(str(tags))


def split_tag_into_subtags(tag):
    print(tag[1:-1])
    if len(tag) >=2 and tag[1].startswith('<'):
        open_tag = tag[0][1:].split()[0]
        close_tag = tag[-1][2:-1]
        if open_tag != close_tag:
            print(tag)
            raise ValueError(f"{open_tag} doesn't equal {close_tag}")
        sub_tags = split_content_into_tags('\n'.join(tag[1:-1]))
        return [tag[0], *[split_tag_into_subtags(this_tag) for this_tag in sub_tags], tag[-1]]
    print(tag)
    return tag

def split_content_into_tags(content):
    # Split content into parts, separating tags and text
    parts = re.split(r'(<[^>]+>|</[^>]+>)', content)
    
    # Remove empty strings from the list
    parts = [part.strip() for part in parts if part.strip()]

    result = []
    current_group = []
    current_tag = None
    opened = 0
    for part in parts:
        if part.startswith('<') and part.endswith('/>'): 
            if len(current_group) == 0:
                result.append([part])
            else:
                current_group.append(part)
            continue

        current_group.append(part)

        if part.startswith('<') and not part.startswith('</'):
            opened += 1
            if current_tag is None:
                current_tag = part

        if current_tag is not None:
            # Closing tag
            if part.startswith('</'):
                opened -= 1
                # Matching closing tag found
                close_tag = part[2:-1]
                open_tag = current_tag[1:].split()[0].replace('>','')
                if close_tag == open_tag and opened == 0:
                    result.append(current_group)
                    current_group = []
                    current_tag = None

    # Add any remaining content
    if len(current_group) > 0:
        result.append(current_group)

    return result

def handle_carousel(doc: DocumentType, tag: str):
    # Extract attributes from the tag
    attributes = dict(re.findall(r'(\w+)="([^"]*)"', tag))
    
    # Example: Read a companion file based on an attribute
    if 'file' in attributes:
        companion_file = attributes['file']
        try:
            with open(companion_file, 'r') as file:
                carousel_content = file.read()
                # Process the carousel content
                doc.add_paragraph(f"Carousel content from {companion_file}:")
                doc.add_paragraph(carousel_content)
        except FileNotFoundError:
            doc.add_paragraph(f"Error: Carousel file {companion_file} not found")
    else:
        doc.add_paragraph("Error: Carousel tag missing 'file' attribute")

# Add more tag handlers as needed

def add_hyperlink(paragraph, url, text):
    # This function adds a hyperlink to a paragraph
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')

    # Add color and underline if needed
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)

    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Join all the xml elements
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)
    return hyperlink


if __name__ == '__main__':
    # Usage
    mdx_to_docx('../stories', 'output.docx')
